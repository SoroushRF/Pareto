import os
import shutil
import json
import time
from typing import List, Optional, Any, Dict, Union
from fastapi import Request
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv
import google.generativeai as genai
from google.generativeai.types import GenerationConfig
from pydantic import BaseModel, Field
import zipfile
import io
from docx import Document
from PIL import Image # Part of Pillow library
from prompts.system_prompt import system_prompt


# 1. SETUP
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")
if not api_key: print("Warning: GEMINI_API_KEY not found.")
genai.configure(api_key=api_key)

MODEL_NAME = "gemini-2.5-flash"
model = genai.GenerativeModel(MODEL_NAME)

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==========================================
# 2. PYDANTIC MODELS (Strict Validation for Critical Data)
# ==========================================

# ==========================================
# 2. UPDATED PYDANTIC MODELS (Lean v2.0)
# ==========================================

class AssessmentAttributes(BaseModel):
    is_bonus: bool = False
    is_mandatory: bool = False
    replacement_logic: Optional[str] = None

class GradingRules(BaseModel):
    # FIX: Optional[int] handles cases where Gemini sends "null" instead of 0
    drop_lowest_n: Optional[int] = 0 
    min_pass_threshold: Optional[float] = None

class TransferPolicy(BaseModel):
    description: Optional[str] = None
    target_id: Optional[str] = None

class AssessmentDate(BaseModel):
    due_date: Optional[str] = None
    is_scheduled_event: bool = False

class AssessmentComponent(BaseModel):
    id: Optional[str] = None
    name: str = "Unknown Assignment"
    
    # Flexible Weight
    weight_percentage: Union[float, str, int] = 0
    quantity: int = 1
    
    # CRITICAL FIX: All nested models must be Optional to handle JSON "null"
    attributes: Optional[AssessmentAttributes] = None
    grading_rules: Optional[GradingRules] = None
    transfer_policy: Optional[TransferPolicy] = None
    dates: Optional[AssessmentDate] = None
    
    evidence: Optional[str] = None

class AssessmentStructure(BaseModel):
    components: List[AssessmentComponent] = []

class GlobalPolicies(BaseModel):
    late_penalty: Optional[str] = "See syllabus"
    missed_work: Optional[str] = "See syllabus"

class OmniscientSyllabus(BaseModel):
    syllabus_metadata: Dict[str, Any] = {}
    course_identity: Dict[str, Any] = {}
    assessment_structure: Optional[AssessmentStructure] = Field(default_factory=AssessmentStructure)
    global_policies: Optional[GlobalPolicies] = Field(default_factory=GlobalPolicies)
    model_config = {"extra": "allow"}
# ==========================================
# 3. THE OMNISCIENT SYSTEM PROMPT (Verbatim 5.1 Template)
# ==========================================
# ==========================================
# 3. LOGIC ADAPTER (Updated for v2.0 - NULL SAFE)
# ==========================================

def organize_syllabus_data(raw_data: dict):
    # 1. Validation
    try:
        syllabus = OmniscientSyllabus(**raw_data)
    except Exception as e:
        print(f"Validation Error: {e}")
        # If validation fails, return an error to UI instead of crashing
        return {
            "error": "Data Validation Failed", 
            "details": str(e),
            "raw_omniscient_json": raw_data
        }

    pareto_assignments = []
    
    components = []
    # Safety check for missing structure
    if syllabus.assessment_structure and syllabus.assessment_structure.components:
         components = syllabus.assessment_structure.components

    for comp in components:
        # --- NULL SAFETY GUARDS ---
        # If Gemini sent 'null', we create empty default objects here
        attrs = comp.attributes or AssessmentAttributes()
        rules = comp.grading_rules or GradingRules()
        transfer = comp.transfer_policy or TransferPolicy()
        dates = comp.dates or AssessmentDate()
        # --------------------------

        category_type = "standard_graded"
        details = {}
        
        # 1. Bonus Check
        if attrs.is_bonus:
            category_type = "standard_graded"
            details["is_bonus"] = True
        
        # 2. Drop Logic (Handle None safely)
        elif (rules.drop_lowest_n or 0) > 0:
            category_type = "internal_drop"
            details = {"drop_count": rules.drop_lowest_n}
            
        # 3. Transfer Logic
        elif transfer.target_id:
            category_type = "external_transfer"
            details = {"transfer_target": transfer.target_id}
            
        # 4. Mandatory Check
        elif attrs.is_mandatory:
            category_type = "strictly_mandatory"
            
        # 5. Variable Weight Check
        if isinstance(comp.weight_percentage, str):
            pass

        pareto_assignments.append({
            "name": comp.name,
            "weight": comp.weight_percentage,
            "type": category_type,
            "details": details,
            "evidence": comp.evidence or "No evidence extracted",
            "due_date": dates.due_date,
            "is_bonus": attrs.is_bonus, 
            "replacement_logic": attrs.replacement_logic 
        })

    # Sort Logic
    def get_sort_weight(item):
        w = item.get('weight', 0)
        return w if isinstance(w, (int, float)) else 0

    pareto_assignments.sort(key=lambda x: (
        1 if x['type'] == 'strictly_mandatory' else 0,
        get_sort_weight(x)
    ), reverse=True)

    return {
        "assignments": pareto_assignments,
        "policies": [
            f"Lateness: {syllabus.global_policies.late_penalty if syllabus.global_policies else 'See Syllabus'}",
            f"Missed Work: {syllabus.global_policies.missed_work if syllabus.global_policies else 'See Syllabus'}"
        ],
        "raw_omniscient_json": raw_data
    }
# ==========================================
# 5. API
# ==========================================

@app.get("/")
def read_root():
    return {
      "status": "Pareto Backend Online",
      "model": MODEL_NAME
    }
# Ensure Request is imported: 
# from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request

def process_docx(file_path):
    """
    Deep Extraction for .docx:
    1. Uses python-docx to read text/tables.
    2. Uses zipfile to extract embedded images (e.g. screenshots of tables).
    """
    content_parts = []
    
    # A. Extract Text
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                # Join cells with pipe to simulate table structure
                row_data = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_data))
        
        text_content = "\n".join(full_text)
        if text_content.strip():
            content_parts.append(text_content)
            print("DEBUG: Extracted text from DOCX.")
    except Exception as e:
        print(f"Warning: Failed to extract text from DOCX: {e}")

    # B. Extract Images (Treating .docx as a ZIP archive)
    try:
        with zipfile.ZipFile(file_path) as z:
            # Look for image files inside the internal structure
            all_files = z.namelist()
            media_files = [f for f in all_files if f.startswith('word/media/')]
            
            for media_path in media_files:
                image_data = z.read(media_path)
                image = Image.open(io.BytesIO(image_data))
                content_parts.append(image)
                print(f"DEBUG: Extracted image: {media_path}")
                
    except Exception as e:
        print(f"Warning: Failed to extract images from DOCX: {e}")
        
    return content_parts

@app.post("/analyze")
async def analyze_syllabus(request: Request, file: UploadFile = File(...)):
    start_time = time.time()
    # Use unique filename to prevent overwrites
    temp_filename = f"temp_{int(time.time())}_{file.filename}" 
    
    try:
        with open(temp_filename, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        gemini_content = [system_prompt]

        if file.filename.lower().endswith(".docx"):
            docx_parts = process_docx(temp_filename)
            gemini_content.extend(docx_parts)
        else:
            uploaded_file = genai.upload_file(temp_filename)
            gemini_content.append(uploaded_file)

        # --- THE FIX IS HERE: FORCE JSON MODE ---
        response_stream = await model.generate_content_async(
            gemini_content, 
            stream=True,
            generation_config=GenerationConfig(
                response_mime_type="application/json"
            )
        )
        # ----------------------------------------

        full_text = ""
        async for chunk in response_stream:
            if await request.is_disconnected():
                os.remove(temp_filename)
                return 
            if chunk.parts:
                full_text += chunk.text
        
        os.remove(temp_filename)
        
        text = full_text
        # Cleaner logic: JSON Mode often removes backticks, so we need to be safe
        if "```json" in text:
            text = text.split("```json")[1].split("```")[0]
        elif "```" in text:
            text = text.split("```")[1].split("```")[0]
        
        # Strip whitespace to prevent empty line errors
        text = text.strip()
            
        raw_data = json.loads(text)
        
        # Add metadata for frontend fallback
        if "syllabus_metadata" not in raw_data:
            raw_data["syllabus_metadata"] = {}
        raw_data["syllabus_metadata"]["source_file_name"] = file.filename

        result = organize_syllabus_data(raw_data)
        result["analysis_duration"] = round(time.time() - start_time, 2)
        
        return result

    except Exception as e:
        print(f"CRITICAL ERROR: {str(e)}")
        # Print the failed text for debugging if it's a JSON error
        if 'text' in locals():
            print(f"FAILED JSON START: {text[:100]}...")
        if os.path.exists(temp_filename):
            os.remove(temp_filename)
        return {"error": str(e)}